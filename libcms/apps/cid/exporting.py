# coding=utf-8
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from django.template import defaultfilters

def _idates_to_word(idates):
    document = Document()

    def style_run(run):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    for idate in sorted(idates, key=lambda x: x.date):
        paragraph = document.add_paragraph('')
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        format = []

        if idate.count_day:
            format.append('j')

        if idate.count_month:
            format.append('E')

        if idate.count_year:
            format.append('Y')

        format = ' '.join(format)

        run = paragraph.add_run((defaultfilters.date(idate.date, format) + ' г').replace(' ', '\u00A0'))
        run.bold = True
        style_run(run)
        run = paragraph.add_run('. ')
        style_run(run)
        run = paragraph.add_run(str(idate).strip().strip('.').strip())
        style_run(run)
        run = paragraph.add_run('.')
        style_run(run)
    doc_file = tempfile.TemporaryFile()
    document.save(doc_file)
    doc_file.seek(0)
    return doc_file