# encoding: utf-8
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from django.db import transaction
from django.http import HttpResponseForbidden
from django.shortcuts import render, HttpResponse, redirect, get_object_or_404
from guardian.decorators import permission_required_or_403
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.contrib.auth.decorators import login_required

from django.template import defaultfilters
from common.pagination import get_page

from .. import models
from ..frontend import views
from forms import TypeForm, ImportantDateForm, FilterForm
from .. import search


@login_required
def index(request):
    if not request.user.has_module_perms('cid'):
        return HttpResponseForbidden()
    return render(request, 'cid/administration/index.html')


@login_required
def id_list(request):
    limit_on_page = 15
    prnt = request.GET.get('print')
    if prnt:
        limit_on_page = 1000
    page = int(request.GET.get('page', 1))
    if page < 1:
        page = 1

    if not request.user.has_module_perms('cid'):
        return HttpResponseForbidden()

    idates_count = models.ImportantDate.objects.all().count()

    attribute = '*'
    value = '*'
    type = ''

    filter_form = FilterForm(request.GET)
    if filter_form.is_valid():
        attribute = filter_form.cleaned_data.get('attribute') or attribute
        value = filter_form.cleaned_data.get('value') or value
        type = filter_form.cleaned_data.get('type') or type

    query = search.construct_query(attr=attribute, value=value, type=type)
    result = search.search(query, start=limit_on_page * (page - 1), rows=limit_on_page, sort=['id_ls desc'])
    paginator = Paginator(result, limit_on_page)

    try:
        paginator_page = paginator.page(page)
    except PageNotAnInteger:
        paginator_page = paginator.page(1)
    except EmptyPage:
        paginator_page = paginator.page(paginator.num_pages)

    ids = [doc['id'] for doc in result.get_docs()]
    idates = search.get_records(ids)
    template = 'cid/administration/id_list.html'
    if prnt:
        if prnt == 'docx':
            with _idates_to_word(idates) as fl:
                response = HttpResponse(fl,
                                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = 'attachment; filename=dates.docx'
                return response
        else:
            template = 'cid/administration/print.html'

    # idates_page = get_page(request, models.ImportantDate.objects.all().order_by('-id'))
    return render(request, template, {
        'idates': idates,
        'num_found': result.get_num_found(),
        'idates_page': paginator_page,
        'idates_count': idates_count,
        'filter_form': filter_form,
    })


@login_required
@permission_required_or_403('cid.add_importantdate')
@transaction.atomic
def create_id(request):
    if request.method == 'POST':
        form = ImportantDateForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('cid:administration:id_list')
    else:
        form = ImportantDateForm()
    return render(request, 'cid/administration/create_id.html', {
        'form': form
    })


@login_required
@transaction.atomic
@permission_required_or_403('cid.change_importantdate')
def edit_id(request, id):
    idate = get_object_or_404(models.ImportantDate, id=id)
    if request.method == 'POST':
        form = ImportantDateForm(request.POST, instance=idate)
        if form.is_valid():
            form.save()
            return redirect('cid:administration:id_list')
    else:
        form = ImportantDateForm(instance=idate)
    return render(request, 'cid/administration/edit_id.html', {
        'form': form
    })


@login_required
@transaction.atomic
@permission_required_or_403('cid.delete_importantdate')
def delete_id(request, id):
    idate = get_object_or_404(models.ImportantDate, id=id)
    idate.delete()
    return redirect('cid:administration:id_list')


def index_important_dates(request):
    models.index_important_dates()
    return redirect('cid:administration:id_list')


def _idates_to_word(idates):
    document = Document()

    def style_run(run):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    for idate in sorted(idates, key=lambda x: x.date):
        paragraph = document.add_paragraph('')
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        format = []

        if idate.count_day:
            format.append('j')

        if idate.count_month:
            format.append('E')

        if idate.count_year:
            format.append('Y')

        format = ' '.join(format)

        run = paragraph.add_run((defaultfilters.date(idate.date, format) + u' г').replace(' ', u'\u00A0'))
        run.bold = True
        style_run(run)
        run = paragraph.add_run(u'. ')
        style_run(run)
        run = paragraph.add_run(unicode(idate).strip().strip('.').strip())
        style_run(run)
        run = paragraph.add_run(u'.')
        style_run(run)
    doc_file = tempfile.TemporaryFile()
    document.save(doc_file)
    doc_file.seek(0)
    return doc_file
